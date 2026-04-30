import math
import random
import re
import sys
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional
from urllib.parse import parse_qs, urlparse

from docx import Document
from selenium import webdriver
from selenium.common.exceptions import (
    NoSuchElementException,
    StaleElementReferenceException,
    TimeoutException,
    WebDriverException,
)
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from webdriver_manager.chrome import ChromeDriverManager


BOARD_URL = "https://cafe.daum.net/school-cont/eri"
OUTPUT_FILE = "학교회계디딤돌_묻고답하기_Re글_수집.docx"
AUTOSAVE_FILE = "autosave.docx"
DOC_TITLE = "학교회계 디딤돌 묻고답하기 Re글 수집"
DEBUG_DIR = Path("debug")
MAX_PAGES_PER_FILE = 300
APPROX_CHARS_PER_LINE = 46
APPROX_LINES_PER_PAGE = 36
LOGIN_NOTICE_HANDLER: Optional[Callable[[WebDriver], None]] = None

WAIT_SECONDS = 15
MIN_ARTICLE_DELAY = 2
MAX_ARTICLE_DELAY = 5
MAX_IFRAME_DEPTH = 5


ARTICLE_TITLE_SELECTORS = [
    "h1",
    "h2",
    "h3",
    ".tit_subject",
    ".tit_info",
    ".article_tit",
    ".article-title",
    ".subject",
    ".title",
    "[class*='tit']",
    "[id*='title']",
]

ARTICLE_BODY_SELECTORS = [
    "#article",
    "#user_contents",
    "#bbs_contents",
    ".article_view",
    ".article-view",
    ".article_content",
    ".article-content",
    ".bbs_contents",
    ".view_content",
    ".view-content",
    ".tx-content-container",
    ".content",
    "[class*='article'][class*='content']",
    "[class*='article'][class*='view']",
    "[id*='content']",
    "[class*='content']",
]

REMOVE_IN_BODY_SELECTORS = [
    "script",
    "style",
    "noscript",
    "iframe",
    "button",
    "form",
    "nav",
    "aside",
    "blockquote",
    ".quote",
    ".quotation",
    ".origin",
    ".original",
    ".comment",
    ".comments",
    ".reply",
    ".replies",
    ".cmt",
    ".memo",
    ".recommend",
    ".vote",
    ".like",
    ".scrap",
    ".share",
    ".attach",
    ".attachment",
    ".file",
    ".tag",
    ".prev",
    ".next",
    ".menu",
    ".banner",
    ".ad",
    ".ads",
    "[id*='comment']",
    "[class*='comment']",
    "[id*='reply']",
    "[class*='reply']",
    "[id*='cmt']",
    "[class*='cmt']",
    "[id*='memo']",
    "[class*='memo']",
    "[class*='recommend']",
    "[class*='attach']",
    "[class*='banner']",
    "[class*='ad']",
]

DROP_BODY_LINE_PATTERNS = [
    re.compile(r"^\s*(댓글|답글|목록|수정|삭제|신고|스크랩|인쇄|추천|공유)\s*$"),
    re.compile(r"^\s*(조회|조회수|추천|추천수)\s*[:：]?\s*\d+\s*$"),
    re.compile(r"^\s*(작성자|작성일|등록일)\s*[:：].*$"),
]

QUOTE_SECTION_PATTERNS = [
    re.compile(r"^\s*-{2,}\s*Original Message\s*-{2,}\s*$", re.I),
    re.compile(r"^\s*=+\s*원문\s*=+\s*$"),
    re.compile(r"^\s*\[?\s*(원문|원본글|이전\s*글)\s*\]?\s*$"),
]

ARTICLE_DATETIME_PATTERNS = [
    re.compile(r"\b(?:\d{2}|\d{4})[./-]\d{1,2}[./-]\d{1,2}\.?\s+\d{1,2}:\d{2}(?::\d{2})?\b"),
    re.compile(r"\b\d{4}년\s*\d{1,2}월\s*\d{1,2}일\s+\d{1,2}:\d{2}(?::\d{2})?\b"),
]


class LoginOrPermissionNoticeError(RuntimeError):
    pass


@dataclass
class ArticleCandidate:
    title: str
    raw_title: str
    key: str
    href: str
    onclick: str


@dataclass
class ArticleData:
    title: str
    published_at: str
    body: str
    key: str


@dataclass
class PageRange:
    start_page: int
    end_page: Optional[int]

    def display(self) -> str:
        end = str(self.end_page) if self.end_page is not None else "끝까지"
        return f"{self.start_page} ~ {end}"


def log(message: str) -> None:
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {message}", flush=True)


def find_cached_chromedriver() -> Optional[str]:
    driver_root = Path.home() / ".wdm" / "drivers" / "chromedriver"
    if not driver_root.exists():
        return None

    candidates = [path for path in driver_root.rglob("chromedriver.exe") if path.is_file()]
    if not candidates:
        return None

    candidates.sort(key=lambda path: path.stat().st_mtime, reverse=True)
    return str(candidates[0])


def setup_driver() -> WebDriver:
    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option("useAutomationExtension", False)

    cached_driver = find_cached_chromedriver()
    if cached_driver:
        try:
            log(f"캐시된 ChromeDriver를 사용합니다: {cached_driver}")
            return webdriver.Chrome(service=Service(cached_driver), options=options)
        except WebDriverException as exc:
            log(f"캐시 ChromeDriver 실행 실패: {exc}")

    try:
        log("Selenium Manager로 ChromeDriver를 준비합니다.")
        return webdriver.Chrome(options=options)
    except WebDriverException as exc:
        log(f"Selenium Manager 실행 실패: {exc}")

    log("webdriver-manager로 ChromeDriver를 준비합니다. 처음 실행이면 시간이 걸릴 수 있습니다.")
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=options)


def random_article_delay() -> None:
    seconds = random.uniform(MIN_ARTICLE_DELAY, MAX_ARTICLE_DELAY)
    log(f"서버 부담을 줄이기 위해 {seconds:.1f}초 대기합니다.")
    time.sleep(seconds)


def ensure_debug_dir() -> None:
    DEBUG_DIR.mkdir(parents=True, exist_ok=True)


def save_debug(driver: WebDriver, label: str) -> None:
    ensure_debug_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_label = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", label).strip("_")[:40] or "debug"
    base = DEBUG_DIR / f"{timestamp}_{safe_label}"

    try:
        driver.save_screenshot(str(base.with_suffix(".png")))
    except WebDriverException as exc:
        log(f"스크린샷 저장 실패: {exc}")

    try:
        source = driver.page_source
        base.with_suffix(".html").write_text(source, encoding="utf-8", errors="replace")
    except WebDriverException as exc:
        log(f"page_source 저장 실패: {exc}")

    log(f"debug 자료 저장: {base.with_suffix('.png').name}, {base.with_suffix('.html').name}")


def compact_spaces(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    return re.sub(r"\s+", " ", text).strip()


def normalize_title(raw_title: str) -> str:
    title = compact_spaces(raw_title)

    # Some boards render reply depth markers or icon alt text before the real title.
    marker_pattern = re.compile(
        r"^\s*(?:"
        r"[ㄴ└┗↳⤷▶▷>+\-·ㆍ•|{}\[\]()]|"
        r"답글|답변|reply|re\.ply|new"
        r")+\s*",
        re.I,
    )

    previous = None
    while previous != title:
        previous = title
        title = marker_pattern.sub("", title).strip()

    return title


def is_re_title(raw_title: str) -> bool:
    normalized = normalize_title(raw_title)
    return bool(re.match(r"^Re\s*:", normalized, re.I))


def cleanup_body_text(text: str, title: str = "") -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines: list[str] = []
    for raw_line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw_line).strip()
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if any(pattern.match(line) for pattern in DROP_BODY_LINE_PATTERNS):
            continue
        lines.append(line)

    while lines and lines[0] == "":
        lines.pop(0)
    while lines and lines[-1] == "":
        lines.pop()

    if title and lines:
        normalized_first = compact_spaces(lines[0])
        normalized_title = compact_spaces(title)
        if normalized_first == normalized_title:
            lines.pop(0)

    cut_at: Optional[int] = None
    for index, line in enumerate(lines):
        if any(pattern.match(line) for pattern in QUOTE_SECTION_PATTERNS):
            cut_at = index
            break
    if cut_at is not None:
        lines = lines[:cut_at]

    return "\n".join(lines).strip()


def readable_text(driver: WebDriver, element: WebElement) -> str:
    try:
        return driver.execute_script(
            "return (arguments[0].innerText || arguments[0].textContent || '').trim();",
            element,
        )
    except WebDriverException:
        try:
            return element.text.strip()
        except WebDriverException:
            return ""


def current_frame_visible_text(driver: WebDriver) -> str:
    try:
        body = driver.find_element(By.TAG_NAME, "body")
    except NoSuchElementException:
        return ""
    return readable_text(driver, body)


def extract_published_at(driver: WebDriver, title: str = "") -> str:
    visible_text = current_frame_visible_text(driver)
    lines = [compact_spaces(line) for line in visible_text.splitlines()]
    lines = [line for line in lines if line]
    if not lines:
        return ""

    search_lines = lines[:25]
    normalized_title = compact_spaces(title)
    if normalized_title:
        for index, line in enumerate(lines):
            if line == normalized_title or normalized_title in line:
                search_lines = lines[index : index + 8]
                break

    for line in search_lines:
        for pattern in ARTICLE_DATETIME_PATTERNS:
            match = pattern.search(line)
            if match:
                return compact_spaces(match.group(0))

    return ""


def is_login_or_permission_notice(text: str) -> bool:
    normalized = compact_spaces(text)
    if not normalized:
        return False

    if "로그인을 하지 않으셨어요" in normalized or "먼저 로그인을" in normalized:
        return True
    if "게시판 권한 안내" in normalized and (
        "로그인하기" in normalized or "카페 가입하기" in normalized
    ):
        return True
    if "읽기권한" in normalized and (
        "읽기가능" in normalized or "카페회원 등급" in normalized
    ):
        return True
    if "접근 권한" in normalized and (
        "없습니다" in normalized or "로그인" in normalized or "가입" in normalized
    ):
        return True
    return False


def has_login_or_permission_notice_in_current_frame(driver: WebDriver) -> bool:
    return is_login_or_permission_notice(current_frame_visible_text(driver))


def element_is_usable(element: WebElement) -> bool:
    try:
        return element.is_displayed() and element.is_enabled()
    except WebDriverException:
        return False


def get_current_frame_url(driver: WebDriver) -> str:
    try:
        return driver.execute_script("return window.location.href;") or ""
    except WebDriverException:
        return ""


def article_key_from_url(url: str, fallback: str = "") -> str:
    if not url:
        return f"title:{compact_spaces(fallback)}"

    parsed = urlparse(url)
    query = parse_qs(parsed.query)
    for name in ("datanum", "dataid", "articleid", "articleId", "articleno", "articleNo", "num", "no"):
        if query.get(name):
            return f"{name}:{query[name][0]}"

    match = re.search(r"(?:bbs_read|article|read)[^0-9]*(\d+)", url, re.I)
    if match:
        return f"article:{match.group(1)}"

    normalized = parsed._replace(fragment="").geturl()
    return normalized or f"title:{compact_spaces(fallback)}"


def element_identity(driver: WebDriver, element: WebElement, title: str) -> str:
    href = element.get_attribute("href") or ""
    onclick = element.get_attribute("onclick") or ""
    if href and not href.lower().startswith("javascript"):
        return article_key_from_url(href, title)
    if onclick:
        number_match = re.search(r"(\d{2,})", onclick)
        if number_match:
            return f"onclick:{number_match.group(1)}"
        return f"onclick:{compact_spaces(onclick)}:{compact_spaces(title)}"
    return f"title:{compact_spaces(title)}"


def get_link_title(driver: WebDriver, element: WebElement) -> str:
    candidates = [
        readable_text(driver, element),
        element.get_attribute("title") or "",
        element.get_attribute("aria-label") or "",
    ]
    for candidate in candidates:
        text = compact_spaces(candidate)
        if text:
            return text
    return ""


def raw_article_link_elements(driver: WebDriver) -> list[WebElement]:
    selectors = [
        "a[href*='bbs_read']",
        "a[href*='article']",
        "a[href*='Article']",
        "a[href*='read']",
        "a.txt_item",
        "a.link_item",
        "td a",
        "li a",
        "a",
    ]
    seen_ids: set[str] = set()
    elements: list[WebElement] = []

    for selector in selectors:
        try:
            found = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in found:
            element_id = getattr(element, "id", "")
            if element_id in seen_ids:
                continue
            seen_ids.add(element_id)
            elements.append(element)
    return elements


def find_re_article_candidates(driver: WebDriver) -> list[ArticleCandidate]:
    candidates: list[ArticleCandidate] = []
    seen_keys: set[str] = set()

    for element in raw_article_link_elements(driver):
        try:
            if not element_is_usable(element):
                continue
            raw_title = get_link_title(driver, element)
            if not raw_title or len(raw_title) > 250:
                continue
            if not is_re_title(raw_title):
                continue

            title = normalize_title(raw_title)
            key = element_identity(driver, element, title)
            if key in seen_keys:
                continue
            seen_keys.add(key)

            candidates.append(
                ArticleCandidate(
                    title=title,
                    raw_title=raw_title,
                    key=key,
                    href=element.get_attribute("href") or "",
                    onclick=element.get_attribute("onclick") or "",
                )
            )
        except StaleElementReferenceException:
            continue
        except WebDriverException:
            continue

    return candidates


def looks_like_article_list(driver: WebDriver) -> bool:
    return len(find_re_article_candidates(driver)) > 0 or any(
        "bbs_read" in (element.get_attribute("href") or "").lower()
        for element in raw_article_link_elements(driver)[:100]
    )


def extract_text_from_clone(driver: WebDriver, element: WebElement) -> str:
    remove_selectors = ",".join(REMOVE_IN_BODY_SELECTORS)
    script = """
        const clone = arguments[0].cloneNode(true);
        const removeSelectors = arguments[1];
        clone.querySelectorAll(removeSelectors).forEach((node) => node.remove());
        return (clone.innerText || clone.textContent || '').trim();
    """
    try:
        return driver.execute_script(script, element, remove_selectors) or ""
    except WebDriverException:
        return readable_text(driver, element)


def extract_title(driver: WebDriver, fallback_title: str = "") -> str:
    best_title = ""
    for selector in ARTICLE_TITLE_SELECTORS:
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            if not element_is_usable(element):
                continue
            text = compact_spaces(readable_text(driver, element))
            if not text or len(text) > 220:
                continue
            if is_re_title(text):
                return normalize_title(text)
            if not best_title and len(text) >= 2:
                best_title = text

    if fallback_title:
        return normalize_title(fallback_title)
    return normalize_title(best_title)


def extract_body(driver: WebDriver, title: str) -> str:
    body_candidates: list[str] = []

    for selector in ARTICLE_BODY_SELECTORS:
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            if not element_is_usable(element):
                continue
            raw_text = extract_text_from_clone(driver, element)
            cleaned = cleanup_body_text(raw_text, title)
            if cleaned:
                body_candidates.append(cleaned)

    if body_candidates:
        body_candidates.sort(key=len, reverse=True)
        return body_candidates[0]

    try:
        body = driver.find_element(By.TAG_NAME, "body")
        return cleanup_body_text(extract_text_from_clone(driver, body), title)
    except NoSuchElementException:
        return ""


def looks_like_article_page(driver: WebDriver) -> bool:
    if has_login_or_permission_notice_in_current_frame(driver):
        return False
    title = extract_title(driver)
    body = extract_body(driver, title)
    return bool(title and body and len(body) >= 5)


def recursive_switch_to_frame(
    driver: WebDriver,
    predicate: Callable[[WebDriver], bool],
    max_depth: int,
    depth: int = 0,
) -> bool:
    try:
        if predicate(driver):
            return True
    except WebDriverException:
        pass

    if depth >= max_depth:
        return False

    try:
        frames = driver.find_elements(By.CSS_SELECTOR, "iframe, frame")
    except WebDriverException:
        return False

    for index in range(len(frames)):
        try:
            frames = driver.find_elements(By.CSS_SELECTOR, "iframe, frame")
            frame = frames[index]
            driver.switch_to.frame(frame)
            if recursive_switch_to_frame(driver, predicate, max_depth, depth + 1):
                return True
            driver.switch_to.parent_frame()
        except (StaleElementReferenceException, WebDriverException):
            try:
                driver.switch_to.parent_frame()
            except WebDriverException:
                driver.switch_to.default_content()
            continue

    return False


def page_has_login_or_permission_notice(driver: WebDriver) -> bool:
    try:
        driver.switch_to.default_content()
        if has_login_or_permission_notice_in_current_frame(driver):
            return True
        return recursive_switch_to_frame(
            driver,
            has_login_or_permission_notice_in_current_frame,
            MAX_IFRAME_DEPTH,
        )
    except WebDriverException:
        return False


def switch_to_matching_frame(
    driver: WebDriver,
    predicate: Callable[[WebDriver], bool],
    description: str,
    timeout: int = WAIT_SECONDS,
) -> None:
    def attempt(current_driver: WebDriver) -> bool:
        current_driver.switch_to.default_content()

        for frame_name in ("down", "cafe_main", "cafeMain", "mainFrame"):
            try:
                current_driver.switch_to.default_content()
                current_driver.switch_to.frame(frame_name)
                if predicate(current_driver):
                    return True
            except WebDriverException:
                continue

        current_driver.switch_to.default_content()
        return recursive_switch_to_frame(current_driver, predicate, MAX_IFRAME_DEPTH)

    WebDriverWait(driver, timeout).until(lambda current_driver: attempt(current_driver))
    log(f"{description} iframe 탐색 완료: {get_current_frame_url(driver)}")


def switch_to_list_frame(driver: WebDriver) -> None:
    switch_to_matching_frame(driver, looks_like_article_list, "게시판 목록")


def switch_to_article_frame(driver: WebDriver) -> None:
    switch_to_matching_frame(driver, looks_like_article_page, "게시글 본문")


def click_re_candidate_by_index(driver: WebDriver, index: int) -> ArticleCandidate:
    elements = raw_article_link_elements(driver)
    filtered: list[tuple[WebElement, ArticleCandidate]] = []
    seen_keys: set[str] = set()

    for element in elements:
        try:
            if not element_is_usable(element):
                continue
            raw_title = get_link_title(driver, element)
            if not raw_title or not is_re_title(raw_title):
                continue
            title = normalize_title(raw_title)
            key = element_identity(driver, element, title)
            if key in seen_keys:
                continue
            seen_keys.add(key)
            filtered.append(
                (
                    element,
                    ArticleCandidate(
                        title=title,
                        raw_title=raw_title,
                        key=key,
                        href=element.get_attribute("href") or "",
                        onclick=element.get_attribute("onclick") or "",
                    ),
                )
            )
        except (StaleElementReferenceException, WebDriverException):
            continue

    if index >= len(filtered):
        raise IndexError("게시글 후보가 새로고침 중 변경되었습니다.")

    element, candidate = filtered[index]
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
    time.sleep(0.3)
    try:
        WebDriverWait(driver, 5).until(EC.element_to_be_clickable(element))
        element.click()
    except WebDriverException:
        driver.execute_script("arguments[0].click();", element)
    return candidate


def make_page_signature(driver: WebDriver) -> str:
    frame_url = get_current_frame_url(driver)
    titles = [candidate.title for candidate in find_re_article_candidates(driver)[:5]]
    return f"{article_key_from_url(frame_url)}|{'|'.join(titles)}"


def get_current_page_number(driver: WebDriver) -> Optional[int]:
    scoped_selectors = [
        ".paging_g .list_paging li.on .num_item",
        ".paging_g .list_paging li.on",
        ".paging_g .on .num_item",
        ".paging_g .on",
        ".pagination .active",
        ".paging .active",
        ".paging .on",
    ]
    for selector in scoped_selectors:
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            if not element_is_usable(element):
                continue
            match = re.search(r"\d{1,5}", compact_spaces(readable_text(driver, element)))
            if match:
                return int(match.group(0))

    frame_url = get_current_frame_url(driver)
    if frame_url:
        page_values = parse_qs(urlparse(frame_url).query).get("page")
        if page_values and page_values[0].isdigit():
            return int(page_values[0])

    try:
        source = driver.page_source
    except WebDriverException:
        source = ""

    for pattern in (
        r"PAGER_page\s*:\s*[\"'](\d+)[\"']",
        r"\bpage\s*:\s*[\"'](\d+)[\"']",
    ):
        match = re.search(pattern, source)
        if match:
            return int(match.group(1))

    return None


def get_last_page_number(driver: WebDriver) -> Optional[int]:
    page_numbers: list[int] = []

    for selector in (
        ".paging_g .list_paging .num_item",
        ".paging_g .list_paging a",
        ".pagination a",
        ".paging a",
    ):
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            if not element_is_usable(element):
                continue
            text = compact_spaces(readable_text(driver, element))
            if re.fullmatch(r"\d{1,5}", text):
                page_numbers.append(int(text))

    if page_numbers:
        return max(page_numbers)

    try:
        source = driver.page_source
    except WebDriverException:
        source = ""

    match = re.search(r"lastPage\s*:\s*[\"'](\d+)[\"']", source)
    if match:
        return int(match.group(1))

    return None


def find_exact_page_link(driver: WebDriver, target_page: int) -> Optional[WebElement]:
    for selector in (
        ".paging_g .list_paging a.link_num",
        ".paging_g .list_paging a",
        ".pagination a",
        ".paging a",
    ):
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            try:
                if not element_is_usable(element):
                    continue
                text = compact_spaces(readable_text(driver, element))
                if re.fullmatch(r"\d{1,5}", text) and int(text) == target_page:
                    return element
            except (StaleElementReferenceException, WebDriverException):
                continue
    return None


def submit_page_form(driver: WebDriver, target_page: int, current_page: Optional[int]) -> bool:
    script = """
        const targetPage = String(arguments[0]);
        const currentPage = arguments[1] === null ? '' : String(arguments[1]);
        const form = document.forms['pageForm'] || document.querySelector('form[name="pageForm"]');
        if (!form || !form.page) {
            return false;
        }
        form.page.value = targetPage;
        if (form.prev_page) {
            form.prev_page.value = currentPage;
        }
        form.submit();
        return true;
    """
    try:
        return bool(driver.execute_script(script, target_page, current_page))
    except WebDriverException:
        return False


def navigate_to_page(driver: WebDriver, target_page: int, current_page: Optional[int]) -> bool:
    before = make_page_signature(driver)
    exact_link = find_exact_page_link(driver, target_page)

    if exact_link is not None:
        click_element(driver, exact_link)
    elif not submit_page_form(driver, target_page, current_page):
        return False

    WebDriverWait(driver, WAIT_SECONDS).until(
        lambda d: get_current_page_number(d) == target_page or make_page_signature(d) != before
    )
    return True


def get_legacy_current_page_number(driver: WebDriver) -> Optional[int]:
    selectors = [
        ".on",
        ".selected",
        ".active",
        ".current",
        "strong",
        "em",
        "span",
    ]
    for selector in selectors:
        try:
            elements = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in elements:
            if not element_is_usable(element):
                continue
            text = compact_spaces(readable_text(driver, element))
            if re.fullmatch(r"\d{1,4}", text):
                return int(text)
    return None


def click_element(driver: WebDriver, element: WebElement) -> None:
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
    time.sleep(0.3)
    try:
        element.click()
    except WebDriverException:
        driver.execute_script("arguments[0].click();", element)


def is_disabled_pagination(element: WebElement) -> bool:
    class_name = (element.get_attribute("class") or "").lower()
    aria_disabled = (element.get_attribute("aria-disabled") or "").lower()
    return "disabled" in class_name or "disable" in class_name or aria_disabled == "true"


def pagination_elements(driver: WebDriver) -> list[WebElement]:
    selectors = [
        ".paging a",
        ".pagination a",
        ".page a",
        "a[class*='page']",
        "a[class*='next']",
        "button[class*='next']",
        "a",
        "button",
    ]
    seen_ids: set[str] = set()
    elements: list[WebElement] = []
    for selector in selectors:
        try:
            found = driver.find_elements(By.CSS_SELECTOR, selector)
        except WebDriverException:
            continue
        for element in found:
            element_id = getattr(element, "id", "")
            if element_id in seen_ids:
                continue
            seen_ids.add(element_id)
            if element_is_usable(element) and not is_disabled_pagination(element):
                elements.append(element)
    return elements


def go_to_next_page(
    driver: WebDriver,
    visited_page_signatures: set[str],
    end_page: Optional[int] = None,
) -> bool:
    current_page = get_current_page_number(driver)
    if current_page is None:
        log("현재 페이지 번호를 찾지 못해 다음 페이지 이동을 중단합니다.")
        save_debug(driver, "current_page_not_found")
        return False

    if end_page is not None and current_page >= end_page:
        log(f"지정한 종료 페이지({end_page})에 도달했습니다.")
        return False

    last_page = get_last_page_number(driver)
    if last_page is not None and current_page >= last_page:
        log("마지막 페이지에 도달했습니다.")
        return False

    target_page = current_page + 1
    if end_page is not None and target_page > end_page:
        log(f"다음 페이지({target_page})가 지정 범위를 넘어 수집을 종료합니다.")
        return False

    try:
        log(f"{current_page} -> {target_page} 페이지로 순차 이동합니다.")
        if not navigate_to_page(driver, target_page, current_page):
            log(f"{target_page} 페이지로 직접 이동하지 못했습니다.")
            save_debug(driver, f"page_move_{target_page}_failed")
            return False

        switch_to_list_frame(driver)
        actual_page = get_current_page_number(driver)
        if actual_page is not None and actual_page != target_page:
            log(f"요청한 페이지는 {target_page}이지만 실제 이동 페이지는 {actual_page}입니다. 순서 보호를 위해 중단합니다.")
            save_debug(driver, f"page_order_mismatch_{target_page}_to_{actual_page}")
            return False

        if make_page_signature(driver) in visited_page_signatures:
            log("이미 방문한 페이지로 이동되어 순회를 종료합니다.")
            return False
        return True
    except (TimeoutException, WebDriverException) as exc:
        log(f"{target_page} 페이지 이동 중 오류: {exc}")
        save_debug(driver, f"page_move_{target_page}_failed")
        return False


def extract_article(driver: WebDriver, fallback_title: str, fallback_key: str) -> ArticleData:
    try:
        switch_to_article_frame(driver)
    except TimeoutException:
        if page_has_login_or_permission_notice(driver):
            save_debug(driver, f"login_or_permission_{fallback_key}")
            raise LoginOrPermissionNoticeError(
                "로그인 또는 게시판 권한 안내문이 표시되어 수집을 멈춥니다."
            )
        raise

    frame_url = get_current_frame_url(driver)
    title = extract_title(driver, fallback_title)
    published_at = extract_published_at(driver, title)
    body = extract_body(driver, title)
    key = article_key_from_url(frame_url, title) if frame_url else fallback_key

    if is_login_or_permission_notice(f"{title}\n{body}"):
        save_debug(driver, f"login_or_permission_{fallback_key}")
        raise LoginOrPermissionNoticeError(
            "로그인 또는 게시판 권한 안내문이 표시되어 수집을 멈춥니다."
        )

    if not title or not body:
        save_debug(driver, f"extract_failed_{fallback_key}")
        raise ValueError("제목 또는 본문 추출에 실패했습니다.")

    return ArticleData(title=title, published_at=published_at, body=body, key=key)


def create_document() -> Document:
    document = Document()
    document.core_properties.title = DOC_TITLE
    document.add_heading(DOC_TITLE, level=0)
    document.add_paragraph(f"수집 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    return document


def estimate_text_lines(text: str) -> int:
    lines = 0
    for raw_line in text.splitlines() or [""]:
        line = raw_line.strip()
        if not line:
            lines += 1
            continue
        lines += max(1, math.ceil(len(line) / APPROX_CHARS_PER_LINE))
    return lines


def estimate_article_pages(article: ArticleData) -> int:
    title_lines = max(2, estimate_text_lines(article.title) + 1)
    published_at_lines = estimate_text_lines(article.published_at) if article.published_at else 0
    body_lines = estimate_text_lines(article.body)
    total_lines = title_lines + published_at_lines + body_lines + 2
    return max(1, math.ceil(total_lines / APPROX_LINES_PER_PAGE))


def numbered_output_file(part_number: int) -> str:
    output_path = Path(OUTPUT_FILE)
    return str(output_path.with_name(f"{output_path.stem}_{part_number:03d}{output_path.suffix}"))


def timestamped_output_file(path: str) -> str:
    original_path = Path(path)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    candidate = original_path.with_name(f"{original_path.stem}_{timestamp}{original_path.suffix}")
    counter = 2

    while candidate.exists():
        candidate = original_path.with_name(
            f"{original_path.stem}_{timestamp}_{counter}{original_path.suffix}"
        )
        counter += 1

    return str(candidate)


class DocumentBatch:
    def __init__(self) -> None:
        self.part_number = 1
        self.document = create_document()
        self.part_article_count = 0
        self.total_collected = 0
        self.estimated_pages = 1
        self.finalized_files: list[str] = []

    @property
    def is_split(self) -> bool:
        return self.part_number > 1 or bool(self.finalized_files)

    def _start_next_part(self) -> None:
        path = numbered_output_file(self.part_number)
        saved_path = save_document(self.document, path)
        self.finalized_files.append(saved_path)

        self.part_number += 1
        self.document = create_document()
        self.part_article_count = 0
        self.estimated_pages = 1
        log(f"{self.part_number}번째 워드파일 작성을 시작합니다.")

    def append(self, article: ArticleData) -> None:
        article_pages = estimate_article_pages(article)
        should_split = (
            self.part_article_count > 0
            and self.estimated_pages + article_pages > MAX_PAGES_PER_FILE
        )
        if should_split:
            log(f"예상 {MAX_PAGES_PER_FILE}페이지에 도달하여 새 워드파일로 분할합니다.")
            self._start_next_part()

        append_article(self.document, article)
        self.part_article_count += 1
        self.total_collected += 1
        self.estimated_pages += article_pages

    def autosave(self) -> None:
        save_document(self.document, AUTOSAVE_FILE)
        if self.is_split:
            save_document(self.document, f"autosave_{self.part_number:03d}.docx")

    def save_final(self) -> list[str]:
        if not self.is_split:
            return [save_document(self.document, OUTPUT_FILE)]

        path = numbered_output_file(self.part_number)
        saved_path = save_document(self.document, path)
        self.finalized_files.append(saved_path)
        return self.finalized_files


def append_article(document: Document, article: ArticleData) -> None:
    document.add_heading(article.title, level=1)
    if article.published_at:
        document.add_paragraph(f"작성일시: {article.published_at}")
    for block in article.body.split("\n\n"):
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        document.add_paragraph("\n".join(lines))


def save_document(document: Document, path: str) -> str:
    try:
        document.save(path)
        log(f"워드파일 저장: {path}")
        return path
    except PermissionError:
        fallback_path = timestamped_output_file(path)
        log(
            f"{path} 파일에 쓸 수 없습니다. Word에서 열려 있거나 Windows가 잠근 상태일 수 있어 "
            f"{fallback_path} 파일로 대신 저장합니다."
        )
        document.save(fallback_path)
        log(f"워드파일 저장: {fallback_path}")
        return fallback_path


def return_to_list(driver: WebDriver) -> None:
    try:
        driver.back()
        switch_to_list_frame(driver)
    except (TimeoutException, WebDriverException):
        log("뒤로가기로 목록 복귀에 실패하여 대상 게시판으로 다시 이동합니다.")
        driver.switch_to.default_content()
        driver.get(BOARD_URL)
        switch_to_list_frame(driver)


def process_current_page(
    driver: WebDriver,
    document_batch: DocumentBatch,
    seen_article_keys: set[str],
) -> None:
    while True:
        switch_to_list_frame(driver)
        candidates = find_re_article_candidates(driver)
        pending_index: Optional[int] = None

        for index, candidate in enumerate(candidates):
            if candidate.key not in seen_article_keys:
                pending_index = index
                break

        if pending_index is None:
            log(f"현재 페이지의 Re: 대상 글 처리 완료: {len(candidates)}개 후보")
            return

        candidate = candidates[pending_index]
        log(f"글 열람: {candidate.title}")

        needs_manual_resume = False
        notice_handler_running = False
        notice_handler_completed = False
        try:
            clicked = click_re_candidate_by_index(driver, pending_index)
            random_article_delay()
            article = extract_article(driver, clicked.title, clicked.key)

            if article.key in seen_article_keys:
                log("이미 수집한 글이라 건너뜁니다.")
            else:
                document_batch.append(article)
                seen_article_keys.add(article.key)
                seen_article_keys.add(clicked.key)
                document_batch.autosave()
                log(
                    f"수집 완료 ({document_batch.total_collected}): "
                    f"{article.title} "
                    f"(현재 파일 예상 {document_batch.estimated_pages}/{MAX_PAGES_PER_FILE}페이지)"
                )
        except LoginOrPermissionNoticeError as exc:
            log(str(exc))
            save_debug(driver, f"login_or_permission_{pending_index}")
            document_batch.autosave()
            if LOGIN_NOTICE_HANDLER is not None:
                notice_handler_running = True
                LOGIN_NOTICE_HANDLER(driver)
                notice_handler_completed = True
                notice_handler_running = False
            else:
                needs_manual_resume = True
                print()
                print("로그인 또는 게시판 권한 안내문이 표시되었습니다.")
                print("열린 Chrome 창에서 직접 로그인하고 게시판 목록이 보이는 상태까지 이동하세요.")
                input("준비가 끝나면 이 콘솔에서 Enter 키를 누르세요: ")
        except Exception as exc:
            log(f"글 처리 중 오류: {exc}")
            save_debug(driver, f"article_error_{pending_index}")
            document_batch.autosave()
            seen_article_keys.add(candidate.key)
        finally:
            if notice_handler_running:
                pass
            elif notice_handler_completed:
                pass
            elif needs_manual_resume:
                driver.switch_to.default_content()
                driver.get(BOARD_URL)
                wait_until_ready_to_collect(driver)
            else:
                return_to_list(driver)


def move_to_start_page(driver: WebDriver, page_range: PageRange) -> None:
    switch_to_list_frame(driver)
    current_page = get_current_page_number(driver)

    if current_page == page_range.start_page:
        log(f"{page_range.start_page} 페이지부터 수집을 시작합니다.")
        return

    log(f"{current_page or '?'} 페이지에서 {page_range.start_page} 페이지로 이동합니다.")
    if not navigate_to_page(driver, page_range.start_page, current_page):
        save_debug(driver, f"start_page_{page_range.start_page}_failed")
        raise RuntimeError(f"{page_range.start_page} 페이지로 이동하지 못했습니다.")

    switch_to_list_frame(driver)
    actual_page = get_current_page_number(driver)
    if actual_page is not None and actual_page != page_range.start_page:
        save_debug(driver, f"start_page_mismatch_{page_range.start_page}_to_{actual_page}")
        raise RuntimeError(
            f"{page_range.start_page} 페이지로 이동하려 했지만 실제 페이지는 {actual_page}입니다."
        )

    log(f"{page_range.start_page} 페이지부터 수집을 시작합니다.")


def crawl_board(driver: WebDriver, page_range: PageRange) -> DocumentBatch:
    document_batch = DocumentBatch()
    seen_article_keys: set[str] = set()
    visited_page_signatures: set[str] = set()

    move_to_start_page(driver, page_range)

    while True:
        switch_to_list_frame(driver)
        current_page = get_current_page_number(driver)
        if (
            page_range.end_page is not None
            and current_page is not None
            and current_page > page_range.end_page
        ):
            log(f"현재 페이지({current_page})가 지정 범위를 넘어 수집을 종료합니다.")
            break

        signature = make_page_signature(driver)
        if signature in visited_page_signatures:
            log("이미 방문한 목록 페이지라 순회를 종료합니다.")
            break
        visited_page_signatures.add(signature)

        process_current_page(driver, document_batch, seen_article_keys)

        switch_to_list_frame(driver)
        current_page = get_current_page_number(driver)
        if (
            page_range.end_page is not None
            and current_page is not None
            and current_page >= page_range.end_page
        ):
            log(f"지정한 종료 페이지({page_range.end_page})까지 처리했습니다.")
            break

        if not go_to_next_page(driver, visited_page_signatures, page_range.end_page):
            log("다음 페이지가 없어 수집을 종료합니다.")
            break

    log(f"총 {document_batch.total_collected}개 글을 수집했습니다.")
    return document_batch


def wait_for_user_ready() -> None:
    print()
    print("크롬 브라우저에서 직접 로그인하고 대상 게시판이 보이는 상태까지 이동하세요.")
    print("로그인 자동화, 캡차 우회, 권한 우회는 하지 않습니다.")
    input("준비가 끝나면 이 콘솔에서 Enter 키를 누르세요: ")


def wait_until_ready_to_collect(driver: WebDriver) -> None:
    while True:
        try:
            driver.switch_to.default_content()

            if page_has_login_or_permission_notice(driver):
                save_debug(driver, "login_or_permission_notice")
                print()
                print("현재 화면은 실제 게시판 목록이 아니라 로그인/권한 안내문입니다.")
                print("열린 Chrome 창에서 직접 로그인하고, 대상 게시판 목록이 보이는 상태까지 이동하세요.")
                input("게시판 목록이 보이면 이 콘솔에서 Enter 키를 누르세요: ")
                continue

            switch_to_list_frame(driver)

            if page_has_login_or_permission_notice(driver):
                save_debug(driver, "login_or_permission_notice")
                print()
                print("현재 화면은 실제 게시판 목록이 아니라 로그인/권한 안내문입니다.")
                print("열린 Chrome 창에서 직접 로그인하고, 대상 게시판 목록이 보이는 상태까지 이동하세요.")
                input("게시판 목록이 보이면 이 콘솔에서 Enter 키를 누르세요: ")
                continue

            return
        except TimeoutException:
            save_debug(driver, "board_list_not_ready")
            print()
            print("게시판 목록을 찾지 못했습니다.")
            print("열린 Chrome 창에서 로그인 상태와 게시판 접근 권한을 확인한 뒤 목록 화면으로 이동하세요.")
            input("게시판 목록이 보이면 이 콘솔에서 Enter 키를 누르세요: ")


def parse_page_number(raw_value: str, field_name: str) -> Optional[int]:
    value = raw_value.strip()
    if not value:
        return None
    if not re.fullmatch(r"\d+", value):
        raise ValueError(f"{field_name}는 숫자만 입력해야 합니다.")
    page_number = int(value)
    if page_number < 1:
        raise ValueError(f"{field_name}는 1 이상이어야 합니다.")
    return page_number


def ask_page_range(driver: WebDriver) -> PageRange:
    switch_to_list_frame(driver)
    current_page = get_current_page_number(driver) or 1
    last_page = get_last_page_number(driver)

    print()
    print("수집할 게시판 페이지 범위를 입력하세요.")
    if last_page is not None:
        print(f"현재 페이지: {current_page}, 마지막 페이지: {last_page}")
    else:
        print(f"현재 페이지: {current_page}, 마지막 페이지: 확인하지 못함")
    print("아무것도 입력하지 않고 Enter를 누르면 기본값을 사용합니다.")

    while True:
        try:
            start_raw = input(f"시작 페이지 (Enter={current_page}): ")
            end_default = str(last_page) if last_page is not None else "끝까지"
            end_raw = input(f"종료 페이지 (Enter={end_default}): ")

            start_page = parse_page_number(start_raw, "시작 페이지") or current_page
            end_page = parse_page_number(end_raw, "종료 페이지")
            if end_page is None:
                end_page = last_page

            if last_page is not None and start_page > last_page:
                raise ValueError(f"시작 페이지가 마지막 페이지({last_page})보다 큽니다.")
            if end_page is not None and start_page > end_page:
                raise ValueError("시작 페이지는 종료 페이지보다 클 수 없습니다.")

            page_range = PageRange(start_page=start_page, end_page=end_page)
            print(f"수집 범위: {page_range.display()} 페이지")
            return page_range
        except ValueError as exc:
            print(f"입력 오류: {exc}")
            print("예: 시작 페이지 10, 종료 페이지 20")
            print()


def main() -> int:
    driver: Optional[WebDriver] = None
    document_batch: Optional[DocumentBatch] = None

    try:
        ensure_debug_dir()
        driver = setup_driver()
        driver.get(BOARD_URL)
        wait_for_user_ready()
        wait_until_ready_to_collect(driver)
        page_range = ask_page_range(driver)

        document_batch = crawl_board(driver, page_range)
        final_files = document_batch.save_final()
        log(f"최종 파일: {', '.join(final_files)}")
        return 0
    except KeyboardInterrupt:
        log("사용자가 중단했습니다.")
        if document_batch is not None:
            document_batch.autosave()
        return 130
    except Exception as exc:
        log(f"예상치 못한 오류: {exc}")
        if driver is not None:
            save_debug(driver, "fatal_error")
        if document_batch is not None:
            document_batch.autosave()
        return 1
    finally:
        if driver is not None:
            print()
            answer = input("브라우저를 닫으려면 Enter 키를 누르세요. 계속 확인하려면 창을 닫지 말고 Ctrl+C로 종료하세요: ")
            _ = answer
            try:
                driver.quit()
            except WebDriverException:
                pass


if __name__ == "__main__":
    sys.exit(main())

